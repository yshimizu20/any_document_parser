from genericpath import isfile
import os
from datetime import datetime
from time import strptime
import csv
from PyPDF2 import PdfReader
import docx


class FileDirectory:
    def __init__(self, dirpath="samples", extracted_dir="samples/extracted/"):
        self.dirpath = dirpath
        self.extracted_dir = extracted_dir
        
    def scrape_title(self):
        """
        scrapes all files in the samples/ directory
        stores all file info in samples/filepaths.csv
        """
        q = [self.dirpath]
        self.titles = []

        while q:
            d = q.pop()
            currdir, nextdir, filedirs = next(os.walk(d))
            for dname in nextdir:
                q.append(d + '/' + dname)
            for filename in filedirs:
                filepath = d + '/' + filename
                date_mod = datetime.fromtimestamp(os.path.getmtime(filepath))
                date_cre = datetime.fromtimestamp(os.path.getctime(filepath))
                parsed_lst = filepath.split(".")
                if len(parsed_lst) == 1:
                    doctype = ""
                else:
                    doctype = parsed_lst[-1]

                self.titles.append({
                    "filepath": filepath,
                    "date_mod": datetime.strftime(date_mod, "%Y-%m-%d"),
                    "date_cre": datetime.strftime(date_cre, "%Y-%m-%d"),
                    "doctype": doctype,
                    "scraped_extracted": "",
                    "scraped_plain": ""
                })
            print(self.titles)

            field_names = ["filepath", "date_mod",
                           "date_cre", "doctype", "scraped_extracted", "scraped_plain"]
            with open("samples/filepaths.csv", "w+") as f:
                writer = csv.DictWriter(f, fieldnames=field_names)
                writer.writeheader()
                writer.writerows(self.titles)

        return self.titles

    def scrape_files(self):
        for file in self.titles:
            parsed = file["filepath"].split('/')[-1].split('.')
            if len(parsed) <= 1:
                continue
            extension = parsed[-1]

            if extension == "txt":
                with open(self.dirpath + file["filepath"], 'r') as f:
                    contents = ''.join(f.readlines()).lower()
                with open(self.extracted_dir + file["filepath"] + '.txt', 'w') as f:
                    f.write(contents)
            
            elif extension == "pdf":
                reader = PdfReader(self.dirpath + file["filepath"])
                contents = ''.join(reader.pages[0].extractText())
                with open(self.extracted_dir + file["filepath"] + '.txt', 'w') as f:
                    f.write(contents)
            
            elif extension == "docx":
                document = docx.Document(self.dirpath + file["filepath"])
                contents = ''.join(document.paragraphs)
                # think about how to process tables
                tables = document.tables

            elif extension == "csv":
                # think about how to process csv tables
                pass
            
            elif extension in ["png", "jpeg", "jpg", "tiff"]:
                # think about how to process images
                pass

            else:
                pass
